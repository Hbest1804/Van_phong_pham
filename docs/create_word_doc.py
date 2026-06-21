"""
Script tạo file Word chứa sơ đồ kiến trúc hệ thống Website Bán Văn Phòng Phẩm
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ─── Page setup: A4 landscape-friendly margins ───
section = doc.sections[0]
section.page_width  = Cm(29.7)
section.page_height = Cm(21.0)
section.left_margin   = Cm(2)
section.right_margin  = Cm(2)
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'),   val.get('val', 'single'))
            border.set(qn('w:sz'),    str(val.get('sz', 6)))
            border.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(border)
    tcPr.append(tcBorders)

def add_heading_paragraph(doc, text, level=1, color_hex=None):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if color_hex:
        for run in p.runs:
            run.font.color.rgb = RGBColor.from_string(color_hex)
    return p

# ══════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════
title = doc.add_heading('SƠ ĐỒ KIẾN TRÚC HỆ THỐNG', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(15, 23, 42)   # dark navy
    run.font.size = Pt(20)

subtitle = doc.add_paragraph('Website Bán Văn Phòng Phẩm Trực Tuyến')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size     = Pt(13)
subtitle.runs[0].font.color.rgb = RGBColor(100, 116, 139)

doc.add_paragraph('')

# ══════════════════════════════════════════════
# ARCHITECTURE IMAGE
# ══════════════════════════════════════════════
img_path = os.path.join(os.path.dirname(__file__), 'system_architecture.png')
if os.path.exists(img_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(9.5))

doc.add_paragraph('')

# ══════════════════════════════════════════════
# SECTION 1 — Mô tả các thành phần
# ══════════════════════════════════════════════
h = doc.add_heading('1. Các thành phần chính', level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(15, 23, 42)

components = [
    ('🌐  Frontend (ReactJS SPA)',  '0E7490', [
        'Framework: ReactJS 19 + Vite 6 + TypeScript',
        'Styling: Tailwind CSS v4',
        'Routing: React Router DOM v7',
        'Port: 3000 (dev) / phục vụ qua Nginx (production)',
        'Tính năng AI: chatbot tư vấn sản phẩm (Gemini)',
    ]),
    ('⚙️  Backend (Node.js / Express)', '6D28D9', [
        'Runtime: Node.js ≥ 18, Framework: Express.js v4',
        'Kiến trúc phân lớp: Routes → Middlewares → Controllers → Services',
        'Authentication: JWT (Access 15m + Refresh 7d, HttpOnly Cookie)',
        'File Upload: Multer → Supabase Storage (giới hạn 5MB)',
        'API prefix: /api/v1  |  Port: 5000',
        'Modules: auth, products, categories, cart, orders, admin, ai, statistics',
    ]),
    ('☁️  Supabase (BaaS)',          '1D4ED8', [
        'Database: PostgreSQL (users, products, categories, cart_items, orders)',
        'Auth: Supabase Auth (tùy chọn — dự án dùng JWT tự quản)',
        'Storage: lưu ảnh sản phẩm, trả về public URL',
        'Client: @supabase/supabase-js (anon key + service_role key)',
    ]),
    ('🤖  Google Gemini AI',         '065F46', [
        'Mô hình: gemini-2.0-flash',
        'Dùng cho: chatbot tư vấn sản phẩm, tìm kiếm ngữ nghĩa',
        'Endpoint backend: POST /api/v1/ai/chat  |  GET /api/v1/ai/search',
        'Lịch sử hội thoại được lưu trong Supabase',
    ]),
    ('📧  SMTP Gmail',               '991B1B', [
        'Thư viện: Nodemailer',
        'Dùng cho: gửi email đặt lại mật khẩu',
        'Token reset hết hạn sau 15 phút (cấu hình được)',
        'Cấu hình: SMTP_HOST, SMTP_USER, SMTP_PASS trong .env',
    ]),
]

for name, color_hex, bullets in components:
    # Component name row
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, color_hex)
    p2 = cell.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r2 = p2.add_run(f'  {name}')
    r2.font.bold  = True
    r2.font.size  = Pt(11)
    r2.font.color.rgb = RGBColor(255, 255, 255)
    cell.add_paragraph('')

    for b in bullets:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.left_indent = Cm(0.5)
        br = bp.add_run(b)
        br.font.size = Pt(10)

    doc.add_paragraph('')

# ══════════════════════════════════════════════
# SECTION 2 — Luồng xử lý Request
# ══════════════════════════════════════════════
h2 = doc.add_heading('2. Luồng xử lý Request', level=1)
for run in h2.runs:
    run.font.color.rgb = RGBColor(15, 23, 42)

flow_steps = [
    ('1', 'Người dùng gửi request từ trình duyệt (ReactJS SPA)',             '06B6D4'),
    ('2', 'Route phân loại endpoint và chuyển đến middleware tương ứng',      '8B5CF6'),
    ('3', 'Middleware xác thực JWT, validate dữ liệu, xử lý file upload',     '8B5CF6'),
    ('4', 'Controller nhận dữ liệu đã xác thực, gọi service',                '8B5CF6'),
    ('5', 'Service thực thi business logic, truy vấn Supabase/Gemini/SMTP',   '8B5CF6'),
    ('6', 'Kết quả được trả về Controller → Response JSON chuẩn → Frontend',  '06B6D4'),
]

tbl2 = doc.add_table(rows=len(flow_steps), cols=2)
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl2.style = 'Table Grid'

for i, (step, desc, color) in enumerate(flow_steps):
    row = tbl2.rows[i]
    # step cell
    c0 = row.cells[0]
    set_cell_bg(c0, color)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(f'Bước {step}')
    r0.font.bold  = True
    r0.font.size  = Pt(10)
    r0.font.color.rgb = RGBColor(255, 255, 255)
    c0.width = Cm(3)

    c1 = row.cells[1]
    p1 = c1.paragraphs[0]
    r1 = p1.add_run(desc)
    r1.font.size = Pt(10)

doc.add_paragraph('')

# ══════════════════════════════════════════════
# SECTION 3 — Bảo mật
# ══════════════════════════════════════════════
h3 = doc.add_heading('3. Bảo mật', level=1)
for run in h3.runs:
    run.font.color.rgb = RGBColor(15, 23, 42)

security_rows = [
    ('Mật khẩu',     'Hash bằng bcryptjs (salt rounds = 12)'),
    ('Access Token', 'JWT — hết hạn sau 15 phút'),
    ('Refresh Token','JWT — hết hạn sau 7 ngày, lưu trong HttpOnly Cookie'),
    ('Token Rotation','Mỗi lần refresh sẽ cấp mới & thu hồi token cũ'),
    ('Phân quyền',   'Middleware authenticate + authorize("admin")'),
    ('Validation',   'express-validator kiểm tra toàn bộ dữ liệu đầu vào'),
    ('File Upload',  'Chỉ chấp nhận image/*, giới hạn 5 MB'),
    ('CORS',         'Chỉ cho phép origin từ domain frontend + *.vercel.app'),
]

tbl3 = doc.add_table(rows=len(security_rows)+1, cols=2)
tbl3.style = 'Table Grid'
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER

# header
for j, hdr in enumerate(['Cơ chế', 'Chi tiết']):
    hc = tbl3.rows[0].cells[j]
    set_cell_bg(hc, '1E293B')
    hp = hc.paragraphs[0]
    hr2 = hp.add_run(hdr)
    hr2.font.bold  = True
    hr2.font.color.rgb = RGBColor(255,255,255)
    hr2.font.size  = Pt(10)

for i, (mech, detail) in enumerate(security_rows, start=1):
    row = tbl3.rows[i]
    mc = row.cells[0]
    set_cell_bg(mc, 'F1F5F9')
    mp = mc.paragraphs[0]
    mr = mp.add_run(mech)
    mr.font.bold = True
    mr.font.size = Pt(10)

    dc = row.cells[1]
    dp = dc.paragraphs[0]
    dr = dp.add_run(detail)
    dr.font.size = Pt(10)

doc.add_paragraph('')

# ══════════════════════════════════════════════
# FOOTER note
# ══════════════════════════════════════════════
note = doc.add_paragraph('📄  Tài liệu được tạo tự động từ mã nguồn dự án  •  Website Bán Văn Phòng Phẩm Trực Tuyến')
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
note.runs[0].font.size = Pt(9)
note.runs[0].font.color.rgb = RGBColor(148, 163, 184)

# ══════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════
output_path = os.path.join(os.path.dirname(__file__), 'Kien_Truc_He_Thong.docx')
doc.save(output_path)
print(f'✅  Đã tạo file Word: {output_path}')
