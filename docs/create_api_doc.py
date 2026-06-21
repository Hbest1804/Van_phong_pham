"""
Script tạo file Word chứa bảng API đầy đủ – Website Bán Văn Phòng Phẩm
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ─── Page setup: A4 landscape ───
section = doc.sections[0]
section.page_width  = Cm(29.7)
section.page_height = Cm(21.0)
section.left_margin   = Cm(1.5)
section.right_margin  = Cm(1.5)
section.top_margin    = Cm(1.5)
section.bottom_margin = Cm(1.5)

# ───────────────────────────── Helpers ──────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_col_width(table, widths_cm):
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            row.cells[i].width = Cm(w)

def method_color(method):
    colors = {
        'GET':    ('0EA5E9', 'FFFFFF'),
        'POST':   ('22C55E', 'FFFFFF'),
        'PUT':    ('F59E0B', '000000'),
        'PATCH':  ('8B5CF6', 'FFFFFF'),
        'DELETE': ('EF4444', 'FFFFFF'),
    }
    return colors.get(method, ('94A3B8', 'FFFFFF'))

def auth_label(auth):
    mapping = {
        '❌': ('E2E8F0', '64748B', 'Public'),
        'User': ('DBEAFE', '1D4ED8', 'User'),
        'Admin': ('FEE2E2', 'DC2626', 'Admin'),
        'User/Guest': ('FEF9C3', 'A16207', 'User/Guest'),
        'User/Admin': ('E0E7FF', '4338CA', 'User/Admin'),
    }
    return mapping.get(auth, ('E2E8F0', '64748B', auth))

def add_group_header(doc, group_title, group_color, base_url, icon=''):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, group_color)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(f'  {icon}  {group_title}')
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(255, 255, 255)

    sub = doc.add_paragraph(f'   Base URL: {base_url}')
    sub.runs[0].font.size = Pt(9)
    sub.runs[0].font.italic = True
    sub.runs[0].font.color.rgb = RGBColor(100, 116, 139)

def add_api_table(doc, endpoints, col_widths):
    """endpoints: list of (method, endpoint, auth, request_params, response, description)"""
    headers = ['Method', 'Endpoint', 'Auth', 'Request Params / Body', 'Response chính', 'Mô tả']
    rows = len(endpoints) + 1
    cols = len(headers)

    t = doc.add_table(rows=rows, cols=cols)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        c = hrow.cells[i]
        set_cell_bg(c, '1E293B')
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for ri, ep in enumerate(endpoints, start=1):
        method, endpoint, auth, req, resp, desc = ep
        row = t.rows[ri]
        bg = 'F8FAFC' if ri % 2 == 0 else 'FFFFFF'

        # METHOD cell
        mc, fc = method_color(method)
        c0 = row.cells[0]
        set_cell_bg(c0, mc)
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(method)
        r0.font.bold = True
        r0.font.size = Pt(9)
        r0.font.color.rgb = RGBColor.from_string(fc)

        # ENDPOINT
        c1 = row.cells[1]
        set_cell_bg(c1, bg)
        r1 = c1.paragraphs[0].add_run(endpoint)
        r1.font.size = Pt(8.5)
        r1.font.bold = True
        r1.font.color.rgb = RGBColor(15, 23, 42)

        # AUTH
        abg, afg, alabel = auth_label(auth)
        c2 = row.cells[2]
        set_cell_bg(c2, abg)
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(alabel)
        r2.font.size = Pt(8.5)
        r2.font.bold = True
        r2.font.color.rgb = RGBColor.from_string(afg)

        # REQUEST
        c3 = row.cells[3]
        set_cell_bg(c3, bg)
        r3 = c3.paragraphs[0].add_run(req)
        r3.font.size = Pt(8)
        r3.font.color.rgb = RGBColor(71, 85, 105)

        # RESPONSE
        c4 = row.cells[4]
        set_cell_bg(c4, bg)
        r4 = c4.paragraphs[0].add_run(resp)
        r4.font.size = Pt(8)
        r4.font.color.rgb = RGBColor(5, 150, 105)

        # DESC
        c5 = row.cells[5]
        set_cell_bg(c5, bg)
        r5 = c5.paragraphs[0].add_run(desc)
        r5.font.size = Pt(8.5)

    set_col_width(t, col_widths)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════
title = doc.add_heading('BẢNG API ENDPOINTS', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(15, 23, 42)
    run.font.size = Pt(22)
    run.font.bold = True

subtitle = doc.add_paragraph('Website Bán Văn Phòng Phẩm Trực Tuyến  —  REST API Documentation')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(12)
subtitle.runs[0].font.color.rgb = RGBColor(100, 116, 139)
subtitle.runs[0].font.bold = True

info = doc.add_paragraph('Base URL: http://localhost:5000   •   Prefix: /api/v1   •   Content-Type: application/json')
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.runs[0].font.size = Pt(9.5)
info.runs[0].font.color.rgb = RGBColor(148, 163, 184)

doc.add_paragraph()

# ── Legend ──────────────────────────────────────────────────────
legend_tbl = doc.add_table(rows=1, cols=6)
legend_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
labels = [
    ('GET',    '0EA5E9', 'FFFFFF'),
    ('POST',   '22C55E', 'FFFFFF'),
    ('PUT',    'F59E0B', '000000'),
    ('PATCH',  '8B5CF6', 'FFFFFF'),
    ('DELETE', 'EF4444', 'FFFFFF'),
    ('LEGEND', '1E293B', 'FFFFFF'),
]
for i, (lbl, bg, fg) in enumerate(labels):
    c = legend_tbl.rows[0].cells[i]
    set_cell_bg(c, bg)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(lbl if lbl != 'LEGEND' else '  HTTP Methods  ')
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(fg)
    c.width = Cm(3.5)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════
# 1. AUTH
# ══════════════════════════════════════════════════════════════════
add_group_header(doc, '1.  Xác thực người dùng  (Auth)', '7C3AED', '/api/v1/auth', '🔐')

auth_endpoints = [
    ('POST',   '/register',        '❌',
     'Body: { email, password, name, phone?, address? }',
     '{ success, data: { user, accessToken } }',
     'Đăng ký tài khoản mới. Mật khẩu ≥6 ký tự, có chữ hoa & số'),
    ('POST',   '/login',           '❌',
     'Body: { email, password }',
     '{ success, data: { user, accessToken } }  +  Set-Cookie: refreshToken',
     'Đăng nhập. Trả về Access Token (15m) và đặt Refresh Token vào HttpOnly Cookie (7d)'),
    ('POST',   '/refresh-token',   '❌',
     'Cookie: refreshToken (HttpOnly tự gửi)',
     '{ success, data: { accessToken } }  +  Set-Cookie: refreshToken mới',
     'Làm mới Access Token. Thực hiện token rotation: thu hồi token cũ, cấp token mới'),
    ('POST',   '/logout',          'User',
     'Cookie: refreshToken (HttpOnly)',
     '{ success, message }  +  Clear-Cookie',
     'Đăng xuất. Thu hồi refresh token, xóa cookie'),
    ('POST',   '/forgot-password', '❌',
     'Body: { email }',
     '{ success, message }  (luôn 200 dù email có tồn tại hay không)',
     'Gửi email chứa link đặt lại mật khẩu. Token hết hạn sau 15 phút'),
    ('POST',   '/reset-password',  '❌',
     'Body: { token, password }',
     '{ success, message }',
     'Đặt lại mật khẩu bằng token từ email. Token chỉ dùng 1 lần'),
]
add_api_table(doc, auth_endpoints, [2.2, 4.5, 2.2, 6.5, 6.0, 5.6])

# ══════════════════════════════════════════════════════════════════
# 2. PRODUCTS
# ══════════════════════════════════════════════════════════════════
add_group_header(doc, '2.  Sản phẩm  (Products)', '16A34A', '/api/v1/products', '📦')

product_endpoints = [
    ('GET',    '/',            '❌',
     'Query: ?page&limit&categoryId&search&minPrice&maxPrice',
     '{ success, data: { products[], total, page, limit, totalPages } }',
     'Danh sách sản phẩm. Hỗ trợ tìm kiếm full-text (pg_trgm), lọc danh mục, phân trang'),
    ('GET',    '/:id',         '❌',
     'Params: id (UUID)',
     '{ success, data: { product } }',
     'Chi tiết một sản phẩm theo ID'),
    ('POST',   '/upload',      'Admin',
     'Form-data: image (file, ≤5MB, image/*)',
     '{ success, data: { imageUrl } }',
     'Upload ảnh lên Supabase Storage, trả về public URL để dùng khi tạo/sửa sản phẩm'),
    ('POST',   '/',            'Admin',
     'Body: { name, description?, price, stock, categoryId, imageUrl? }',
     '{ success, data: { product } }',
     'Tạo sản phẩm mới. price ≥ 0, stock ≥ 0, categoryId là UUID hợp lệ'),
    ('PUT',    '/:id',         'Admin',
     'Params: id  |  Body: { name?, description?, price?, stock?, categoryId?, imageUrl?, isActive? }',
     '{ success, data: { product } }',
     'Cập nhật sản phẩm (partial update). Tất cả field đều optional'),
    ('DELETE', '/:id',         'Admin',
     'Params: id (UUID)',
     '{ success, message }',
     'Xóa sản phẩm. Sẽ lỗi nếu sản phẩm đang có trong đơn hàng (RESTRICT)'),
]
add_api_table(doc, product_endpoints, [2.2, 4.5, 2.2, 6.5, 6.0, 5.6])

# ══════════════════════════════════════════════════════════════════
# 3. CATEGORIES
# ══════════════════════════════════════════════════════════════════
add_group_header(doc, '3.  Danh mục  (Categories)', '0E7490', '/api/v1/categories', '🏷️')

cat_endpoints = [
    ('GET',    '/',    '❌',
     '—',
     '{ success, data: { categories[] } }',
     'Lấy toàn bộ danh mục sản phẩm'),
    ('GET',    '/:id', '❌',
     'Params: id (UUID)',
     '{ success, data: { category } }',
     'Chi tiết một danh mục'),
    ('POST',   '/',    'Admin',
     'Body: { name, description? }',
     '{ success, data: { category } }',
     'Tạo danh mục mới. Tên danh mục phải duy nhất'),
    ('PUT',    '/:id', 'Admin',
     'Params: id  |  Body: { name?, description? }',
     '{ success, data: { category } }',
     'Cập nhật danh mục'),
    ('DELETE', '/:id', 'Admin',
     'Params: id (UUID)',
     '{ success, message }',
     'Xóa danh mục. Sẽ lỗi nếu còn sản phẩm thuộc danh mục này (RESTRICT)'),
]
add_api_table(doc, cat_endpoints, [2.2, 4.5, 2.2, 6.5, 6.0, 5.6])

# ══════════════════════════════════════════════════════════════════
# 4. CART
# ══════════════════════════════════════════════════════════════════
add_group_header(doc, '4.  Giỏ hàng  (Cart)', '0D9488', '/api/v1/cart', '🛒')

cart_endpoints = [
    ('GET',    '/',              'User',
     '—',
     '{ success, data: { items[], totalItems, totalPrice } }',
     'Lấy toàn bộ giỏ hàng của user hiện tại (server-side cart)'),
    ('POST',   '/',              'User',
     'Body: { productId (UUID), quantity? (default: 1) }',
     '{ success, data: { cartItem } }',
     'Thêm sản phẩm vào giỏ. Nếu đã có → tự tăng số lượng. Kiểm tra tồn kho atomic'),
    ('POST',   '/bulk-sync',     'User',
     'Body: { items: [{ productId, quantity }] }',
     '{ success, message }',
     'Đồng bộ giỏ hàng local/guest khi đăng nhập. Merge với giỏ server, giới hạn theo tồn kho'),
    ('PUT',    '/:productId',    'User',
     'Params: productId (UUID)  |  Body: { quantity (≥1) }',
     '{ success, data: { cartItem } }',
     'Cập nhật số lượng sản phẩm trong giỏ'),
    ('DELETE', '/:productId',    'User',
     'Params: productId (UUID)',
     '{ success, message }',
     'Xóa một sản phẩm khỏi giỏ hàng'),
    ('DELETE', '/',              'User',
     '—',
     '{ success, message }',
     'Xóa toàn bộ giỏ hàng'),
]
add_api_table(doc, cart_endpoints, [2.2, 4.5, 2.2, 6.5, 6.0, 5.6])

# ══════════════════════════════════════════════════════════════════
# 5. ORDERS
# ══════════════════════════════════════════════════════════════════
add_group_header(doc, '5.  Đơn hàng  (Orders)', 'D97706', '/api/v1/orders', '🧾')

order_endpoints = [
    ('POST',  '/',             'User',
     'Body: { address, paymentMethod? (cod|transfer), note? }',
     '{ success, data: { orderId } }',
     'Tạo đơn hàng từ giỏ hàng hiện tại (atomic transaction). Validate tồn kho, giảm stock, xóa cart'),
    ('GET',   '/my',           'User',
     'Query: ?page (≥1)&limit (1–50)',
     '{ success, data: { orders[], total, page, limit, totalPages } }',
     'Lấy danh sách đơn hàng của chính user đang đăng nhập, kèm phân trang'),
    ('GET',   '/:id',          'User/Admin',
     'Params: id (UUID)',
     '{ success, data: { order, items[] } }',
     'Chi tiết đơn hàng. User chỉ xem được đơn của mình, Admin xem được tất cả'),
    ('GET',   '/',             'Admin',
     'Query: ?page&limit&status (pending|shipping|completed|cancelled)',
     '{ success, data: { orders[], total, page, limit, totalPages } }',
     'Lấy tất cả đơn hàng với bộ lọc theo trạng thái và phân trang — chỉ Admin'),
    ('PATCH', '/:id/status',   'Admin',
     'Params: id  |  Body: { status (pending|shipping|completed|cancelled) }',
     '{ success, data: { order } }',
     'Cập nhật trạng thái đơn hàng. Trigger tự hoàn kho khi cancelled'),
]
add_api_table(doc, order_endpoints, [2.2, 4.5, 2.2, 6.5, 6.0, 5.6])

# ══════════════════════════════════════════════════════════════════
# 6. ADMIN USERS
# ══════════════════════════════════════════════════════════════════
add_group_header(doc, '6.  Quản lý người dùng  (Admin Users)', '4F46E5', '/api/v1/admin/users', '👥')

admin_user_endpoints = [
    ('GET',   '/',            'Admin',
     'Query: ?page&limit&status (active|locked)&sortBy&sortOrder',
     '{ success, data: { users[], total, page, limit, totalPages } }',
     'Danh sách tất cả người dùng với bộ lọc trạng thái, sắp xếp và phân trang'),
    ('GET',   '/:id',         'Admin',
     'Params: id (UUID)',
     '{ success, data: { user } }',
     'Xem thông tin chi tiết một người dùng'),
    ('PATCH', '/:id/status',  'Admin',
     'Params: id  |  Body: { status (active|locked) }',
     '{ success, data: { user } }',
     'Khóa hoặc mở khóa tài khoản người dùng'),
]
add_api_table(doc, admin_user_endpoints, [2.2, 4.5, 2.2, 6.5, 6.0, 5.6])

# ══════════════════════════════════════════════════════════════════
# 7. STATISTICS
# ══════════════════════════════════════════════════════════════════
add_group_header(doc, '7.  Thống kê  (Statistics)', '9D174D', '/api/v1/admin/statistics', '📊')

stats_endpoints = [
    ('GET', '/overview',     'Admin',
     '—',
     '{ summary, currentMonth, previousMonth, ordersByStatus, topProducts[] }',
     'Tổng quan Dashboard: doanh thu, so sánh tháng trước, đơn theo trạng thái, top 5 sản phẩm'),
    ('GET', '/revenue',      'Admin',
     'Query: ?from (ISO 8601)&to (ISO 8601)&groupBy (month|day)',
     '{ period, groupBy, data: [{ label, revenue, orders }], summary }',
     'Doanh thu theo khoảng thời gian, nhóm theo ngày hoặc tháng'),
    ('GET', '/top-products', 'Admin',
     'Query: ?limit (1–50)&from&to&sortBy (quantity|revenue)',
     '{ period, sortBy, limit, products: [{ rank, productName, unitPrice, totalQuantity, totalRevenue }], summary }',
     'Top sản phẩm bán chạy nhất, sắp xếp theo số lượng hoặc doanh thu'),
]
add_api_table(doc, stats_endpoints, [2.2, 4.5, 2.2, 6.5, 6.0, 5.6])

# ══════════════════════════════════════════════════════════════════
# 8. AI ADVISOR
# ══════════════════════════════════════════════════════════════════
add_group_header(doc, '8.  Tư vấn AI  (AI Advisor)', '065F46', '/api/v1/ai', '🤖')

ai_endpoints = [
    ('POST',   '/chat',                  'User/Guest',
     'Body: { message, sessionId? }',
     '{ success, data: { reply, sessionId, sessionTitle } }',
     'Gửi tin nhắn đến AI Gemini. Lưu lịch sử hội thoại trong Supabase'),
    ('GET',    '/search',                'User/Guest',
     'Query: ?q (chuỗi tìm kiếm ngôn ngữ tự nhiên)',
     '{ success, data: { products[] } }',
     'Tìm kiếm sản phẩm bằng ngôn ngữ tự nhiên thông qua AI (semantic search)'),
    ('GET',    '/sessions',              'User/Guest',
     '—',
     '{ success, data: { sessions[] } }',
     'Lấy danh sách tất cả phiên hội thoại AI của user / guest hiện tại'),
    ('GET',    '/sessions/:session_id',  'User/Guest',
     'Params: session_id (UUID)',
     '{ success, data: { session, messages[] } }',
     'Lấy toàn bộ tin nhắn trong một phiên hội thoại cụ thể'),
    ('DELETE', '/sessions/:session_id',  'User/Guest',
     'Params: session_id (UUID)',
     '{ success, message }',
     'Xóa phiên hội thoại và tất cả tin nhắn liên quan'),
]
add_api_table(doc, ai_endpoints, [2.2, 4.5, 2.2, 6.5, 6.0, 5.6])

# ══════════════════════════════════════════════════════════════════
# 9. HEALTH CHECK
# ══════════════════════════════════════════════════════════════════
add_group_header(doc, '9.  Health Check', '475569', '/api/health', '🏥')

health_endpoints = [
    ('GET', '/api/health', '❌',
     '—',
     '{ status: "ok", timestamp }',
     'Kiểm tra trạng thái hoạt động của server. Dùng cho Docker HEALTHCHECK'),
]
add_api_table(doc, health_endpoints, [2.2, 4.5, 2.2, 6.5, 6.0, 5.6])

# ══════════════════════════════════════════════════════════════════
# PHẦN PHỤ LỤC — Auth & Error codes
# ══════════════════════════════════════════════════════════════════
doc.add_paragraph()
h = doc.add_heading('Phụ lục — Cấu trúc Response & Mã lỗi', level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(15, 23, 42)

# Response structure
res_info = [
    ('Success', '{ "success": true, "data": { ... } }', '200 / 201'),
    ('Error',   '{ "success": false, "message": "...", "errors"?: [...] }', '400 / 401 / 403 / 404 / 500'),
]

tbl_res = doc.add_table(rows=len(res_info)+1, cols=3)
tbl_res.style = 'Table Grid'
tbl_res.alignment = WD_TABLE_ALIGNMENT.LEFT

for i, h in enumerate(['Loại', 'Cấu trúc JSON', 'HTTP Status']):
    c = tbl_res.rows[0].cells[i]
    set_cell_bg(c, '334155')
    p = c.paragraphs[0]
    r = p.add_run(h)
    r.font.bold = True; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(255,255,255)

for ri, (typ, struct, status) in enumerate(res_info, start=1):
    row = tbl_res.rows[ri]
    for ci, txt in enumerate([typ, struct, status]):
        c = row.cells[ci]
        set_cell_bg(c, 'F8FAFC')
        r = c.paragraphs[0].add_run(txt)
        r.font.size = Pt(9)
        r.font.bold = (ci == 0)

for row in tbl_res.rows:
    row.cells[0].width = Cm(3)
    row.cells[1].width = Cm(16)
    row.cells[2].width = Cm(7)

doc.add_paragraph()

# HTTP Status codes
status_codes = [
    ('200 OK',                'Thành công'),
    ('201 Created',           'Tạo tài nguyên thành công'),
    ('400 Bad Request',       'Dữ liệu đầu vào không hợp lệ (validation error)'),
    ('401 Unauthorized',      'Chưa đăng nhập hoặc token hết hạn'),
    ('403 Forbidden',         'Không có quyền truy cập (sai role)'),
    ('404 Not Found',         'Tài nguyên không tồn tại'),
    ('409 Conflict',          'Trùng dữ liệu (email đã tồn tại, tên danh mục trùng...)'),
    ('500 Internal Error',    'Lỗi máy chủ nội bộ'),
]

sc_tbl = doc.add_table(rows=len(status_codes)+1, cols=2)
sc_tbl.style = 'Table Grid'
sc_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

for i, h in enumerate(['HTTP Status Code', 'Ý nghĩa']):
    c = sc_tbl.rows[0].cells[i]
    set_cell_bg(c, '334155')
    r = c.paragraphs[0].add_run(h)
    r.font.bold = True; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(255,255,255)

for ri, (code, meaning) in enumerate(status_codes, start=1):
    row = sc_tbl.rows[ri]
    bg = 'FEF2F2' if code.startswith(('4','5')) else 'F0FDF4'
    c0, c1 = row.cells[0], row.cells[1]
    set_cell_bg(c0, bg); set_cell_bg(c1, bg)
    r0 = c0.paragraphs[0].add_run(code)
    r0.font.size = Pt(9); r0.font.bold = True
    r0.font.color.rgb = RGBColor(220,38,38) if code.startswith(('4','5')) else RGBColor(5,150,105)
    r1 = c1.paragraphs[0].add_run(meaning)
    r1.font.size = Pt(9)
    c0.width = Cm(5); c1.width = Cm(21)

doc.add_paragraph()

# Footer
note = doc.add_paragraph(
    '📄  Tài liệu tạo tự động từ mã nguồn dự án  •  Website Bán Văn Phòng Phẩm Trực Tuyến  •  Node.js / Express.js / PostgreSQL'
)
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
note.runs[0].font.size = Pt(8)
note.runs[0].font.color.rgb = RGBColor(148, 163, 184)

# ══════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════
output_path = os.path.join(os.path.dirname(__file__), 'API_Documentation.docx')
doc.save(output_path)
print(f'✅  Đã tạo file Word: {output_path}')
