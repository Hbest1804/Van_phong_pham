"""
Script tạo file Word chứa Sơ đồ ERD – Database Website Bán Văn Phòng Phẩm
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ─── Page setup: A4 portrait ───
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin   = Cm(1.5)
section.right_margin  = Cm(1.5)
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)

# ─── Helper functions ───────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color='0F172A', align=WD_ALIGN_PARAGRAPH.LEFT):
    h = doc.add_heading(text, level=level)
    h.alignment = align
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return h

def add_table_row(table, row_idx, col0_text, col1_text, col0_bold=False,
                  col0_bg=None, col1_bg=None, font_size=9):
    row = table.rows[row_idx]
    c0, c1 = row.cells[0], row.cells[1]
    if col0_bg: set_cell_bg(c0, col0_bg)
    if col1_bg: set_cell_bg(c1, col1_bg)
    r0 = c0.paragraphs[0].add_run(col0_text)
    r0.font.size = Pt(font_size)
    r0.font.bold = col0_bold
    r1 = c1.paragraphs[0].add_run(col1_text)
    r1.font.size = Pt(font_size)

def make_header_row(table, headers, bg_color, font_color='FFFFFF'):
    row = table.rows[0]
    for i, hdr in enumerate(headers):
        cell = row.cells[i]
        set_cell_bg(cell, bg_color)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(hdr)
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(font_color)

# ══════════════════════════════════════════════════════════════════
# TRANG BÌA
# ══════════════════════════════════════════════════════════════════
title = doc.add_heading('SƠ ĐỒ THỰC THỂ - QUAN HỆ (ERD)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(15, 23, 42)
    run.font.size = Pt(22)
    run.font.bold = True

subtitle = doc.add_paragraph('Database – Website Bán Văn Phòng Phẩm Trực Tuyến')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(100, 116, 139)
subtitle.runs[0].font.bold = True

meta = doc.add_paragraph('PostgreSQL (Supabase)  •  10 bảng  •  4 ENUM types  •  7 Stored Functions')
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.runs[0].font.size = Pt(10)
meta.runs[0].font.color.rgb = RGBColor(148, 163, 184)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════
# SƠ ĐỒ ERD (Hình ảnh)
# ══════════════════════════════════════════════════════════════════
add_heading(doc, '1. Sơ Đồ ERD Tổng Quan', level=1)

img_path = os.path.join(os.path.dirname(__file__), 'erd_database.png')
if os.path.exists(img_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(6.5))
else:
    doc.add_paragraph('[⚠️ Không tìm thấy file erd_database.png]')

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════
# PHẦN 2 — MÔ TẢ BẢNG
# ══════════════════════════════════════════════════════════════════
add_heading(doc, '2. Mô Tả Chi Tiết Các Bảng', level=1)

# ─── Danh sách bảng ───
tables_info = [
    {
        'name': 'users',
        'desc': 'Tài khoản người dùng (khách hàng & quản trị viên)',
        'color': '7C3AED',
        'columns': [
            ('id',            'UUID',           'PK',  'Khóa chính, tự sinh'),
            ('email',         'VARCHAR(255)',    'UK',  'Email đăng nhập, duy nhất'),
            ('password_hash', 'TEXT',            '',    'Mật khẩu băm bcryptjs (salt 12)'),
            ('name',          'VARCHAR(100)',    '',    'Tên hiển thị'),
            ('role',          'ENUM',            '',    "user | admin"),
            ('status',        'ENUM',            '',    "active | locked"),
            ('phone',         'VARCHAR(20)',     '',    'Số điện thoại (tuỳ chọn)'),
            ('address',       'TEXT',            '',    'Địa chỉ mặc định'),
            ('created_at',    'TIMESTAMPTZ',     '',    'Thời điểm tạo'),
            ('updated_at',    'TIMESTAMPTZ',     '',    'Tự cập nhật qua trigger'),
        ]
    },
    {
        'name': 'refresh_tokens',
        'desc': 'Lưu Refresh Token để hỗ trợ đăng xuất & xoay vòng token',
        'color': '1D4ED8',
        'columns': [
            ('id',         'UUID',        'PK', 'Khóa chính'),
            ('user_id',    'UUID',        'FK', 'Tham chiếu users.id → CASCADE'),
            ('token_hash', 'TEXT',        'UK', 'Hash của refresh token'),
            ('expires_at', 'TIMESTAMPTZ', '',   'Thời hạn hết hiệu lực'),
            ('revoked',    'BOOLEAN',     '',   'TRUE = đã thu hồi'),
            ('created_at', 'TIMESTAMPTZ', '',   'Thời điểm cấp token'),
        ]
    },
    {
        'name': 'password_reset_tokens',
        'desc': 'Token đặt lại mật khẩu (tối đa 1 token/user, hết hạn 15 phút)',
        'color': '0E7490',
        'columns': [
            ('id',         'UUID',        'PK',    'Khóa chính'),
            ('user_id',    'UUID',        'FK/UK', 'Tham chiếu users.id, UNIQUE'),
            ('token_hash', 'TEXT',        'UK',    'SHA-256 hash gửi qua email'),
            ('expires_at', 'TIMESTAMPTZ', '',      'Thời hạn hết hiệu lực'),
            ('created_at', 'TIMESTAMPTZ', '',      'Thời điểm tạo'),
        ]
    },
    {
        'name': 'categories',
        'desc': 'Danh mục sản phẩm văn phòng phẩm',
        'color': '059669',
        'columns': [
            ('id',          'UUID',         'PK', 'Khóa chính'),
            ('name',        'VARCHAR(100)', 'UK', 'Tên danh mục, duy nhất'),
            ('description', 'TEXT',         '',   'Mô tả danh mục'),
            ('created_at',  'TIMESTAMPTZ',  '',   'Thời điểm tạo'),
            ('updated_at',  'TIMESTAMPTZ',  '',   'Tự cập nhật qua trigger'),
        ]
    },
    {
        'name': 'products',
        'desc': 'Sản phẩm văn phòng phẩm (bút, giấy, file, băng keo…)',
        'color': '16A34A',
        'columns': [
            ('id',          'UUID',          'PK', 'Khóa chính'),
            ('name',        'VARCHAR(255)',   '',   'Tên sản phẩm'),
            ('description', 'TEXT',          '',   'Mô tả chi tiết'),
            ('price',       'NUMERIC(12,0)', '',   'Giá bán VNĐ (≥ 0)'),
            ('stock',       'INTEGER',       '',   'Tồn kho (≥ 0), trigger tự giảm'),
            ('category_id', 'UUID',          'FK', 'Tham chiếu categories.id → RESTRICT'),
            ('image_url',   'TEXT',          '',   'URL ảnh trên Supabase Storage'),
            ('is_active',   'BOOLEAN',       '',   'FALSE = ẩn sản phẩm'),
            ('created_at',  'TIMESTAMPTZ',   '',   'Thời điểm tạo'),
            ('updated_at',  'TIMESTAMPTZ',   '',   'Tự cập nhật qua trigger'),
        ]
    },
    {
        'name': 'orders',
        'desc': 'Đơn hàng của khách hàng',
        'color': 'D97706',
        'columns': [
            ('id',             'UUID',          'PK', 'Khóa chính'),
            ('user_id',        'UUID',          'FK', 'Tham chiếu users.id → RESTRICT'),
            ('status',         'ENUM',          '',   'pending | shipping | completed | cancelled'),
            ('total',          'NUMERIC(15,0)', '',   'Tổng tiền VNĐ (snapshot)'),
            ('address',        'TEXT',          '',   'Địa chỉ giao hàng (snapshot)'),
            ('payment_method', 'ENUM',          '',   'cod | transfer'),
            ('note',           'TEXT',          '',   'Ghi chú của khách'),
            ('created_at',     'TIMESTAMPTZ',   '',   'Thời điểm đặt hàng'),
            ('updated_at',     'TIMESTAMPTZ',   '',   'Tự cập nhật qua trigger'),
        ]
    },
    {
        'name': 'order_items',
        'desc': 'Chi tiết từng mặt hàng trong đơn hàng (snapshot giá & tên)',
        'color': 'EA580C',
        'columns': [
            ('id',           'UUID',          'PK', 'Khóa chính'),
            ('order_id',     'UUID',          'FK', 'Tham chiếu orders.id → CASCADE'),
            ('product_id',   'UUID',          'FK', 'Tham chiếu products.id → RESTRICT'),
            ('product_name', 'VARCHAR(255)',   '',   'Snapshot tên sản phẩm lúc đặt'),
            ('quantity',     'INTEGER',       '',   'Số lượng (> 0)'),
            ('unit_price',   'NUMERIC(12,0)', '',   'Snapshot giá lúc đặt (VNĐ)'),
        ]
    },
    {
        'name': 'cart_items',
        'desc': 'Giỏ hàng phía server — mỗi row là 1 sản phẩm trong giỏ',
        'color': '0D9488',
        'columns': [
            ('id',         'UUID',        'PK', 'Khóa chính'),
            ('user_id',    'UUID',        'FK', 'Tham chiếu users.id → CASCADE'),
            ('product_id', 'UUID',        'FK', 'Tham chiếu products.id → CASCADE'),
            ('quantity',   'INTEGER',     '',   'Số lượng (> 0, DEFAULT 1)'),
            ('created_at', 'TIMESTAMPTZ', '',   'Thời điểm thêm vào giỏ'),
            ('updated_at', 'TIMESTAMPTZ', '',   'Tự cập nhật qua trigger'),
        ]
    },
    {
        'name': 'chat_sessions',
        'desc': 'Phiên trò chuyện với AI chatbot',
        'color': 'DB2777',
        'columns': [
            ('id',               'UUID',        'PK', 'Khóa chính'),
            ('user_id',          'UUID',        'FK', 'Tham chiếu users.id (nullable — guest OK)'),
            ('guest_session_id', 'UUID',        '',   'ID phiên khách (không đăng nhập)'),
            ('title',            'VARCHAR(255)', '',   "Tiêu đề phiên chat"),
            ('created_at',       'TIMESTAMPTZ', '',   'Thời điểm tạo'),
            ('updated_at',       'TIMESTAMPTZ', '',   'Tự cập nhật qua trigger'),
        ]
    },
    {
        'name': 'chat_messages',
        'desc': 'Tin nhắn trong phiên trò chuyện AI',
        'color': 'BE185D',
        'columns': [
            ('id',         'UUID',        'PK', 'Khóa chính'),
            ('session_id', 'UUID',        'FK', 'Tham chiếu chat_sessions.id → CASCADE'),
            ('sender',     'VARCHAR(10)', '',   "user | ai"),
            ('message',    'TEXT',        '',   'Nội dung tin nhắn'),
            ('created_at', 'TIMESTAMPTZ', '',   'Thời điểm gửi'),
        ]
    },
]

for tbl in tables_info:
    # Tên bảng
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cell = t.cell(0, 0)
    set_cell_bg(hdr_cell, tbl['color'])
    hp = hdr_cell.paragraphs[0]
    hr = hp.add_run(f"  📋  {tbl['name'].upper()}")
    hr.font.bold = True
    hr.font.size = Pt(11)
    hr.font.color.rgb = RGBColor(255, 255, 255)

    # Mô tả
    desc_p = doc.add_paragraph(f"   {tbl['desc']}")
    desc_p.runs[0].font.size = Pt(9)
    desc_p.runs[0].font.italic = True
    desc_p.runs[0].font.color.rgb = RGBColor(100, 116, 139)

    # Bảng cột
    col_count = 4
    col_table = doc.add_table(rows=len(tbl['columns']) + 1, cols=col_count)
    col_table.style = 'Table Grid'
    col_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header
    make_header_row(col_table, ['Cột', 'Kiểu dữ liệu', 'Ràng buộc', 'Ghi chú'], '1E293B')

    # Rows
    for i, (col_name, dtype, constraint, note) in enumerate(tbl['columns'], start=1):
        row = col_table.rows[i]
        bg = 'F8FAFC' if i % 2 == 0 else 'FFFFFF'

        # Col name
        c0 = row.cells[0]
        set_cell_bg(c0, bg)
        r0 = c0.paragraphs[0].add_run(col_name)
        r0.font.size = Pt(9)
        r0.font.bold = True
        if constraint == 'PK':
            r0.font.color.rgb = RGBColor(220, 38, 38)
        elif 'FK' in constraint:
            r0.font.color.rgb = RGBColor(37, 99, 235)

        # Datatype
        c1 = row.cells[1]
        set_cell_bg(c1, bg)
        r1 = c1.paragraphs[0].add_run(dtype)
        r1.font.size = Pt(9)
        r1.font.color.rgb = RGBColor(5, 150, 105)

        # Constraint
        c2 = row.cells[2]
        set_cell_bg(c2, bg)
        r2 = c2.paragraphs[0].add_run(constraint)
        r2.font.size = Pt(9)
        r2.font.bold = True
        if constraint == 'PK':
            r2.font.color.rgb = RGBColor(220, 38, 38)
        elif 'FK' in constraint:
            r2.font.color.rgb = RGBColor(37, 99, 235)
        elif 'UK' in constraint:
            r2.font.color.rgb = RGBColor(217, 119, 6)

        # Note
        c3 = row.cells[3]
        set_cell_bg(c3, bg)
        r3 = c3.paragraphs[0].add_run(note)
        r3.font.size = Pt(9)
        r3.font.color.rgb = RGBColor(71, 85, 105)

    # Cột width
    for row in col_table.rows:
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(3.5)
        row.cells[2].width = Cm(2.5)
        row.cells[3].width = Cm(8.5)

    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════
# PHẦN 3 — QUAN HỆ GIỮA CÁC BẢNG
# ══════════════════════════════════════════════════════════════════
add_heading(doc, '3. Các Mối Quan Hệ (Relationships)', level=1)

relations = [
    ('users → refresh_tokens',        '1 → N', 'ON DELETE CASCADE',  'Xoá user → xoá toàn bộ token'),
    ('users → password_reset_tokens', '1 → 1', 'ON DELETE CASCADE',  'Mỗi user tối đa 1 reset token'),
    ('users → orders',                '1 → N', 'ON DELETE RESTRICT', 'Xoá user → lỗi nếu còn đơn hàng'),
    ('users → cart_items',            '1 → N', 'ON DELETE CASCADE',  'Xoá user → xoá toàn bộ giỏ hàng'),
    ('users → chat_sessions',         '1 → N', 'nullable FK',        'Nullable — guest không cần tài khoản'),
    ('categories → products',         '1 → N', 'ON DELETE RESTRICT', 'Xoá category → lỗi nếu còn sản phẩm'),
    ('products → order_items',        '1 → N', 'ON DELETE RESTRICT', 'Xoá product → lỗi nếu đã có trong đơn'),
    ('products → cart_items',         '1 → N', 'ON DELETE CASCADE',  'Xoá product → xoá khỏi giỏ hàng'),
    ('orders → order_items',          '1 → N', 'ON DELETE CASCADE',  'Xoá order → xoá toàn bộ chi tiết'),
    ('chat_sessions → chat_messages', '1 → N', 'ON DELETE CASCADE',  'Xoá phiên → xoá toàn bộ tin nhắn'),
]

rel_table = doc.add_table(rows=len(relations) + 1, cols=4)
rel_table.style = 'Table Grid'
rel_table.alignment = WD_TABLE_ALIGNMENT.CENTER

make_header_row(rel_table, ['Quan hệ', 'Loại', 'Referential Action', 'Ghi chú'], '334155')

for i, (rel, kind, action, note) in enumerate(relations, start=1):
    row = rel_table.rows[i]
    bg = 'F8FAFC' if i % 2 == 0 else 'FFFFFF'
    for j, (text, bold) in enumerate([(rel, True), (kind, True), (action, False), (note, False)]):
        c = row.cells[j]
        set_cell_bg(c, bg)
        r = c.paragraphs[0].add_run(text)
        r.font.size = Pt(9)
        r.font.bold = bold
        if j == 1:
            r.font.color.rgb = RGBColor(124, 58, 237)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════
# PHẦN 4 — ENUM TYPES
# ══════════════════════════════════════════════════════════════════
add_heading(doc, '4. ENUM Types', level=1)

enums = [
    ('user_role',      'users.role',          'user, admin'),
    ('user_status',    'users.status',         'active, locked'),
    ('order_status',   'orders.status',        'pending, shipping, completed, cancelled'),
    ('payment_method', 'orders.payment_method','cod, transfer'),
]

enum_table = doc.add_table(rows=len(enums) + 1, cols=3)
enum_table.style = 'Table Grid'
make_header_row(enum_table, ['ENUM Name', 'Dùng ở cột', 'Các giá trị'], '334155')

for i, (name, col, vals) in enumerate(enums, start=1):
    row = enum_table.rows[i]
    bg = 'F8FAFC' if i % 2 == 0 else 'FFFFFF'
    for j, text in enumerate([name, col, vals]):
        c = row.cells[j]
        set_cell_bg(c, bg)
        r = c.paragraphs[0].add_run(text)
        r.font.size = Pt(9)
        r.font.bold = (j == 0)
        if j == 2:
            r.font.color.rgb = RGBColor(5, 150, 105)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════
# PHẦN 5 — TRIGGERS & STORED FUNCTIONS
# ══════════════════════════════════════════════════════════════════
add_heading(doc, '5. Triggers', level=1)

triggers = [
    ('trg_*_updated_at',                    'BEFORE UPDATE',         'users, categories, products, orders, cart_items, chat_sessions',    'Tự cập nhật updated_at = NOW()'),
    ('trg_order_items_decrease_stock',      'AFTER INSERT',          'order_items',    'Giảm products.stock theo quantity đặt hàng'),
    ('trg_orders_manage_stock',             'BEFORE UPDATE (status)','orders',         'Hoàn kho khi → cancelled; chặn đảo ngược từ cancelled'),
    ('trg_orders_restore_stock_on_delete',  'BEFORE DELETE',         'orders',         'Hoàn kho khi xoá đơn chưa huỷ'),
]

trg_table = doc.add_table(rows=len(triggers) + 1, cols=4)
trg_table.style = 'Table Grid'
make_header_row(trg_table, ['Trigger', 'Sự kiện', 'Bảng', 'Hành động'], '334155')

for i, (name, event, table, action) in enumerate(triggers, start=1):
    row = trg_table.rows[i]
    bg = 'F8FAFC' if i % 2 == 0 else 'FFFFFF'
    for j, text in enumerate([name, event, table, action]):
        c = row.cells[j]
        set_cell_bg(c, bg)
        r = c.paragraphs[0].add_run(text)
        r.font.size = Pt(9)
        r.font.bold = (j == 0)
        if j == 1:
            r.font.color.rgb = RGBColor(220, 38, 38)

doc.add_paragraph()

add_heading(doc, '6. Stored Functions', level=1)

funcs = [
    ('add_to_cart_atomic()',        'Thêm/cập nhật giỏ hàng atomic — kiểm tra tồn kho & trạng thái sản phẩm'),
    ('sync_cart_bulk()',            'Đồng bộ giỏ hàng guest → server khi đăng nhập (JSONB input)'),
    ('create_order_transaction()',  'Tạo đơn hàng atomic: validate → insert orders → insert items → clear cart'),
    ('get_revenue_sum()',           'Tổng doanh thu đơn completed trong khoảng thời gian'),
    ('get_top_products()',          'Danh sách sản phẩm bán chạy nhất (sort by quantity/revenue)'),
    ('get_top_products_summary()',  'Tổng hợp số liệu top sản phẩm'),
    ('get_revenue_stats()',         'Thống kê doanh thu & đơn hàng theo ngày hoặc tháng'),
]

fn_table = doc.add_table(rows=len(funcs) + 1, cols=2)
fn_table.style = 'Table Grid'
make_header_row(fn_table, ['Hàm', 'Mục đích'], '334155')

for i, (name, desc) in enumerate(funcs, start=1):
    row = fn_table.rows[i]
    bg = 'F8FAFC' if i % 2 == 0 else 'FFFFFF'
    c0, c1 = row.cells[0], row.cells[1]
    set_cell_bg(c0, bg); set_cell_bg(c1, bg)
    r0 = c0.paragraphs[0].add_run(name)
    r0.font.size = Pt(9); r0.font.bold = True
    r0.font.color.rgb = RGBColor(124, 58, 237)
    r1 = c1.paragraphs[0].add_run(desc)
    r1.font.size = Pt(9)
    c0.width = Cm(5.5); c1.width = Cm(13.0)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════
note = doc.add_paragraph(
    '📄  Tài liệu được tạo tự động từ schema.sql  •  Website Bán Văn Phòng Phẩm Trực Tuyến  •  PostgreSQL / Supabase'
)
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
note.runs[0].font.size = Pt(8)
note.runs[0].font.color.rgb = RGBColor(148, 163, 184)

# ══════════════════════════════════════════════════════════════════
# LƯU FILE
# ══════════════════════════════════════════════════════════════════
output_path = os.path.join(os.path.dirname(__file__), 'ERD_Database.docx')
doc.save(output_path)
print(f'✅  Đã tạo file Word: {output_path}')
